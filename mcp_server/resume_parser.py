#!/usr/bin/env python3
"""Resume file parser supporting .doc, .docx, .pdf, .txt formats.

Extracts text from various resume file formats and cleans non-UTF8 characters.
"""

from __future__ import annotations

import os
import re
from pathlib import Path


def clean_text(text: str) -> str:
    """Remove non-UTF8 characters and normalize whitespace."""
    # Remove non-printable characters except newlines, tabs, carriage returns
    text = re.sub(r'[^\x20-\x7E\n\r\t\u0080-\uFFFF]+', ' ', text)

    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n+', '\n\n', text)

    # Strip leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    return text.strip()


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from .txt file with encoding fallback."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                return clean_text(f.read())
        except (UnicodeDecodeError, LookupError):
            continue

    # Last resort: binary read and decode with errors ignored
    with open(file_path, 'rb') as f:
        raw = f.read()
        text = raw.decode('utf-8', errors='ignore')
        return clean_text(text)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from .docx file using python-docx."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx files. Install with: pip install python-docx"
        )

    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        text = '\n'.join(paragraphs)
        return clean_text(text)
    except Exception as e:
        raise ValueError(f"Failed to parse .docx file: {e}")


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from .pdf file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except ImportError:
        raise ImportError(
            "PyPDF2 is required for .pdf files. Install with: pip install PyPDF2"
        )

    try:
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = '\n\n'.join(text_parts)
        return clean_text(text)
    except Exception as e:
        raise ValueError(f"Failed to parse .pdf file: {e}")


def extract_text_from_doc(file_path: str) -> str:
    """Extract text from legacy .doc file.

    Note: .doc parsing is complex and requires external tools like antiword or LibreOffice.
    This implementation attempts basic text extraction but may not work for all .doc files.
    """
    # For .doc files, we'll try to read as binary and extract text patterns
    # This is a fallback and may not work well for all .doc files
    try:
        with open(file_path, 'rb') as f:
            raw = f.read()
            # Try to decode as latin-1 first (common in older Word docs)
            try:
                text = raw.decode('latin-1', errors='ignore')
            except UnicodeDecodeError:
                text = raw.decode('utf-8', errors='ignore')

            # Remove control characters and binary junk
            text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)
            return clean_text(text)
    except Exception as e:
        raise ValueError(
            f"Failed to parse .doc file: {e}. "
            "Consider converting to .docx or .pdf for better results."
        )


def parse_resume_file(file_path: str) -> str:
    """Parse resume file and return cleaned text content.

    Supports: .txt, .docx, .pdf, .doc

    Args:
        file_path: Path to resume file

    Returns:
        Cleaned text content from the resume

    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file format is unsupported or parsing fails
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == '.txt':
        return extract_text_from_txt(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.doc':
        return extract_text_from_doc(file_path)
    else:
        raise ValueError(
            f"Unsupported file format: {ext}. "
            "Supported formats: .txt, .docx, .pdf, .doc"
        )


__all__ = ['parse_resume_file', 'clean_text']
